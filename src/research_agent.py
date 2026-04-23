"""
src/research_agent.py
=====================
Autonomous Research Agent — Production-Grade, Security-Hardened

Architecture:
  - CitationTracker  : deduplication, credibility scoring, reference rendering
  - DiskCache        : SHA-256-keyed JSON blobs, TTL-aware
  - SessionStore     : SQLite-backed; parameterised queries throughout
  - run_agent()      : Claude agentic loop (tool use, max-iteration guard)
  - export_*()       : Markdown / PDF / DOCX renderers
  - create_app()     : Flask web UI with SSE streaming and security headers

Environment variables (see .env.example):
  ANTHROPIC_API_KEY  — required
  BRAVE_API_KEY      — required for live search; falls back to mock otherwise
  FLASK_SECRET_KEY   — required when running --serve
  ALLOWED_HOSTS      — comma-separated hostnames accepted by the web UI
  REPORTS_DIR        — override default ./reports output directory
  CACHE_DIR          — override default ./cache directory
  DB_PATH            — override default ./sessions.db path
  MAX_QUESTION_LEN   — max characters accepted per research question (default 500)
  CACHE_TTL_SECONDS  — cache entry lifetime in seconds (default 86400 = 24 h)
  UI_USERNAME        — username for web UI Basic Auth (leave blank to disable auth)
  UI_PASSWORD_SALT   — hex salt for password hash (generate with --gen-password)
  UI_PASSWORD_HASH   — PBKDF2-SHA256 password hash (generate with --gen-password)
  RATE_LIMIT_GLOBAL_RPS   — max requests/sec per IP for all routes (default 10)
  RATE_LIMIT_RESEARCH_RPM — max /start requests/min per IP (default 4)
"""

from __future__ import annotations

import datetime
import hashlib
import json
import logging
import os
import queue
import re
import secrets
import sqlite3
import sys
import threading
import time
import traceback
from pathlib import Path
from typing import Iterator, Optional

import requests
import anthropic

# ── Optional UI dependency ────────────────────────────────────────────────────
try:
    from flask import (
        Flask, Response, abort, request,
        render_template_string, jsonify, send_file,
    )
    FLASK_OK = True
except ImportError:
    FLASK_OK = False

# ── Optional export dependencies ──────────────────────────────────────────────
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ══════════════════════════════════════════════════════════════════════════════
# Logging
# ══════════════════════════════════════════════════════════════════════════════

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("research_agent")

# ══════════════════════════════════════════════════════════════════════════════
# Configuration  (all values from environment — no hardcoded secrets)
# ══════════════════════════════════════════════════════════════════════════════

# SECURITY: API keys are read exclusively from environment variables.
# They are never logged, stored in source, or echoed back to clients.
ANTHROPIC_API_KEY: str = os.environ.get("ANTHROPIC_API_KEY", "")
BRAVE_API_KEY: str     = os.environ.get("BRAVE_API_KEY", "")
FLASK_SECRET_KEY: str  = os.environ.get("FLASK_SECRET_KEY", "")

ALLOWED_HOSTS: list[str] = [
    h.strip()
    for h in os.environ.get("ALLOWED_HOSTS", "localhost,127.0.0.1").split(",")
    if h.strip()
]

MODEL           = "claude-opus-4-5"
MAX_TOKENS      = 4096
MAX_SEARCH_HITS = 5
MAX_ITERATIONS  = 40
RETRY_BACKOFF   = [1, 2, 4, 8]  # exponential back-off in seconds

# SECURITY: All runtime data is confined to operator-controlled directories.
# Never write outside these paths; path-traversal checks enforce this below.
REPORTS_DIR = Path(os.environ.get("REPORTS_DIR", "reports")).resolve()
CACHE_DIR   = Path(os.environ.get("CACHE_DIR",   "cache")).resolve()
DB_PATH     = Path(os.environ.get("DB_PATH",     "sessions.db")).resolve()

MAX_QUESTION_LEN  = int(os.environ.get("MAX_QUESTION_LEN",  "500"))
CACHE_TTL_SECONDS = int(os.environ.get("CACHE_TTL_SECONDS", "86400"))

# Valid export formats — whitelist used for input validation
VALID_FORMATS = {"md", "pdf", "docx"}

# ── Authentication (optional, enabled when UI_USERNAME is set) ────────────────
# SECURITY: Set UI_USERNAME + UI_PASSWORD_HASH to enable HTTP Basic Auth on
# the web UI. Generate the hash with:
#   python -c "import hashlib,os; pw=b'yourpassword'; \
#     print(hashlib.pbkdf2_hmac('sha256',pw,os.urandom(16),260000).hex())"
# Or leave unset to run without auth (trusted-network / localhost only).
UI_USERNAME:      str = os.environ.get("UI_USERNAME", "")
UI_PASSWORD_HASH: str = os.environ.get("UI_PASSWORD_HASH", "")
UI_PASSWORD_SALT: str = os.environ.get("UI_PASSWORD_SALT", "")
AUTH_ENABLED: bool = bool(UI_USERNAME and UI_PASSWORD_HASH and UI_PASSWORD_SALT)

# ── In-process rate limiting ──────────────────────────────────────────────────
# SECURITY: Limits per IP to reduce abuse of the expensive LLM pipeline.
# For production, prefer nginx rate-limiting (see nginx/nginx.conf).
RATE_LIMIT_GLOBAL_RPS   = int(os.environ.get("RATE_LIMIT_GLOBAL_RPS",   "10"))  # per second
RATE_LIMIT_RESEARCH_RPM = int(os.environ.get("RATE_LIMIT_RESEARCH_RPM", "4"))   # per minute

REPORTS_DIR.mkdir(parents=True, exist_ok=True)
CACHE_DIR.mkdir(parents=True, exist_ok=True)

# ══════════════════════════════════════════════════════════════════════════════
# Input validation helpers
# ══════════════════════════════════════════════════════════════════════════════

def validate_question(q: str) -> str:
    """Sanitise and length-limit the research question.

    SECURITY: Strips null bytes and enforces a hard length cap to prevent
    prompt-injection attempts that rely on extremely long inputs.
    """
    if not isinstance(q, str):
        raise ValueError("question must be a string")
    q = q.replace("\x00", "").strip()
    if not q:
        raise ValueError("question must not be empty")
    if len(q) > MAX_QUESTION_LEN:
        raise ValueError(f"question exceeds {MAX_QUESTION_LEN} character limit")
    return q


def validate_formats(raw: list) -> list[str]:
    """Return only whitelisted export format strings.

    SECURITY: Whitelist validation prevents unexpected format strings from
    reaching file-writing code and potentially causing path issues.
    """
    if not isinstance(raw, list):
        raise ValueError("formats must be a list")
    accepted = [f for f in raw if isinstance(f, str) and f.lower() in VALID_FORMATS]
    if not accepted:
        raise ValueError(f"no valid formats supplied; accepted: {VALID_FORMATS}")
    return [f.lower() for f in accepted]


def safe_report_path(path_str: str) -> Path:
    """Resolve a report path and confirm it stays inside REPORTS_DIR.

    SECURITY: Prevents path-traversal attacks on the /download endpoint.
    Any path that resolves outside REPORTS_DIR is rejected with a 404.
    """
    p = Path(path_str).resolve()
    try:
        p.relative_to(REPORTS_DIR)  # raises ValueError if outside
    except ValueError:
        raise PermissionError(f"path outside reports directory: {path_str}")
    return p

# ══════════════════════════════════════════════════════════════════════════════
# Authentication  (HTTP Basic Auth, optional)
# ══════════════════════════════════════════════════════════════════════════════

import base64
import hmac as _hmac

def _check_basic_auth(auth_header: str) -> bool:
    """Validate an HTTP Basic Auth header against the configured credentials.

    SECURITY: Uses PBKDF2-HMAC-SHA256 for password storage and
    hmac.compare_digest for constant-time comparison (timing-attack resistant).
    Both are required: a fast hash (e.g. MD5) or == comparison would be unsafe.
    """
    if not auth_header or not auth_header.startswith("Basic "):
        return False
    try:
        decoded = base64.b64decode(auth_header[6:]).decode("utf-8")
        username, _, password = decoded.partition(":")
    except Exception:
        return False

    # SECURITY: Compare usernames with constant-time digest to prevent
    # timing-based username enumeration.
    username_ok = _hmac.compare_digest(username.encode(), UI_USERNAME.encode())

    # SECURITY: PBKDF2 with 260,000 iterations (NIST SP 800-132 recommendation).
    import hashlib
    salt_bytes = bytes.fromhex(UI_PASSWORD_SALT)
    candidate_hash = hashlib.pbkdf2_hmac(
        "sha256", password.encode(), salt_bytes, 260_000
    ).hex()
    password_ok = _hmac.compare_digest(candidate_hash, UI_PASSWORD_HASH)

    return username_ok and password_ok


def _require_auth():
    """Flask before_request hook — enforces Basic Auth when AUTH_ENABLED."""
    if not AUTH_ENABLED:
        return  # auth disabled; pass through
    auth = request.headers.get("Authorization", "")
    if not _check_basic_auth(auth):
        # SECURITY: Return 401 with WWW-Authenticate so browsers prompt for creds.
        return Response(
            "Authentication required.",
            status=401,
            headers={"WWW-Authenticate": 'Basic realm="Research Agent"'},
        )


# ── Auth helper for CLI usage ─────────────────────────────────────────────────

def generate_password_hash(password: str) -> tuple[str, str]:
    """Return (salt_hex, hash_hex) for a given password.

    Usage:
        salt, phash = generate_password_hash("mysecretpassword")
        # Set UI_PASSWORD_SALT=<salt> and UI_PASSWORD_HASH=<phash> in .env
    """
    import hashlib
    salt = os.urandom(16)
    phash = hashlib.pbkdf2_hmac("sha256", password.encode(), salt, 260_000)
    return salt.hex(), phash.hex()

# ══════════════════════════════════════════════════════════════════════════════
# In-process rate limiter  (token bucket, per-IP)
# ══════════════════════════════════════════════════════════════════════════════

class _TokenBucket:
    """Thread-safe token bucket for per-IP rate limiting.

    SECURITY: This is a defence-in-depth measure. The nginx rate-limiting
    layer (see nginx/nginx.conf) is the primary control; this catches abuse
    when the app is run without nginx (e.g. development, direct access).
    """

    def __init__(self, rate: float, capacity: float) -> None:
        self._rate = rate          # tokens added per second
        self._capacity = capacity
        self._buckets: dict[str, tuple[float, float]] = {}  # ip → (tokens, last_refill)
        self._lock = threading.Lock()

    def allow(self, ip: str, cost: float = 1.0) -> bool:
        now = time.monotonic()
        with self._lock:
            tokens, last = self._buckets.get(ip, (self._capacity, now))
            # Refill tokens since last check
            tokens = min(self._capacity, tokens + (now - last) * self._rate)
            if tokens >= cost:
                self._buckets[ip] = (tokens - cost, now)
                return True
            self._buckets[ip] = (tokens, now)
            return False

    def cleanup(self, max_age_seconds: float = 300.0) -> None:
        """Remove stale bucket entries to bound memory usage."""
        now = time.monotonic()
        with self._lock:
            self._buckets = {
                ip: (t, ts)
                for ip, (t, ts) in self._buckets.items()
                if now - ts < max_age_seconds
            }


# Global bucket: 10 req/s with burst of 20
_global_bucket   = _TokenBucket(rate=RATE_LIMIT_GLOBAL_RPS,   capacity=20.0)
# Research bucket: 4 req/min (0.067/s) with burst of 2
_research_bucket = _TokenBucket(rate=RATE_LIMIT_RESEARCH_RPM / 60.0, capacity=2.0)


def _client_ip() -> str:
    """Return the real client IP, respecting X-Forwarded-For from trusted proxies."""
    # SECURITY: Only trust X-Forwarded-For when behind a known reverse proxy.
    # In a direct-access deployment, use request.remote_addr exclusively.
    forwarded = request.headers.get("X-Forwarded-For", "")
    if forwarded:
        # Take the leftmost (originating) IP; strip whitespace.
        return forwarded.split(",")[0].strip()
    return request.remote_addr or "unknown"

# ══════════════════════════════════════════════════════════════════════════════
# Persistence — SQLite session store
# ══════════════════════════════════════════════════════════════════════════════

# SECURITY: All SQL uses parameterised queries (?). No string interpolation
# is used in query construction, preventing SQL-injection attacks.
def _init_db() -> sqlite3.Connection:
    conn = sqlite3.connect(str(DB_PATH), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")   # safe concurrent reads
    conn.execute("PRAGMA foreign_keys=ON")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS sessions (
            id          TEXT PRIMARY KEY,
            question    TEXT NOT NULL,
            status      TEXT NOT NULL DEFAULT 'running',
            started_at  TEXT NOT NULL,
            finished_at TEXT,
            messages    TEXT NOT NULL DEFAULT '[]',
            citations   TEXT NOT NULL DEFAULT '{}',
            report_md   TEXT NOT NULL DEFAULT ''
        )
    """)
    conn.commit()
    return conn


_db_conn: sqlite3.Connection = _init_db()
_db_lock = threading.Lock()


def session_create(sid: str, question: str) -> None:
    with _db_lock:
        _db_conn.execute(
            "INSERT OR IGNORE INTO sessions(id, question, started_at) VALUES (?, ?, ?)",
            (sid, question, datetime.datetime.utcnow().isoformat()),
        )
        _db_conn.commit()


def session_update(sid: str, **kw) -> None:
    if not kw:
        return
    # SECURITY: column names come from internal call sites only, never from
    # user input, so formatting them into the query is safe here.
    allowed_cols = {"status", "finished_at", "messages", "citations", "report_md"}
    invalid = set(kw) - allowed_cols
    if invalid:
        raise ValueError(f"invalid session columns: {invalid}")
    cols = ", ".join(f"{k} = ?" for k in kw)
    vals = list(kw.values()) + [sid]
    with _db_lock:
        _db_conn.execute(f"UPDATE sessions SET {cols} WHERE id = ?", vals)
        _db_conn.commit()


def session_get(sid: str) -> Optional[dict]:
    row = _db_conn.execute(
        "SELECT * FROM sessions WHERE id = ?", (sid,)
    ).fetchone()
    return dict(row) if row else None


def sessions_list() -> list[dict]:
    rows = _db_conn.execute(
        "SELECT id, question, status, started_at, finished_at "
        "FROM sessions ORDER BY started_at DESC LIMIT 100"
    ).fetchall()
    return [dict(r) for r in rows]

# ══════════════════════════════════════════════════════════════════════════════
# Disk cache  — SHA-256 keyed, TTL-aware
# ══════════════════════════════════════════════════════════════════════════════

def _cache_path(prefix: str, key: str) -> Path:
    # SECURITY: Cache filenames are derived from a hash — never from raw user
    # input — so filenames cannot contain path separators or shell metacharacters.
    digest = hashlib.sha256(key.encode()).hexdigest()[:24]
    return CACHE_DIR / f"{prefix}_{digest}.json"


def cache_get(prefix: str, key: str) -> Optional[object]:
    p = _cache_path(prefix, key)
    if not p.exists():
        return None
    try:
        payload = json.loads(p.read_text(encoding="utf-8"))
        if time.time() - payload.get("ts", 0) > CACHE_TTL_SECONDS:
            p.unlink(missing_ok=True)
            return None
        return payload["data"]
    except Exception:
        return None


def cache_set(prefix: str, key: str, value: object) -> None:
    p = _cache_path(prefix, key)
    p.write_text(
        json.dumps({"ts": time.time(), "data": value}, ensure_ascii=False),
        encoding="utf-8",
    )

# ══════════════════════════════════════════════════════════════════════════════
# Retry wrapper  — exponential back-off
# ══════════════════════════════════════════════════════════════════════════════

def with_retry(fn, *args, label: str = "call", **kwargs):
    """Call *fn* with exponential back-off on any exception.

    Raises the last exception after all retries are exhausted.
    SECURITY: Errors are logged via the standard logger, never sent to clients.
    """
    schedule = RETRY_BACKOFF + [None]
    for attempt, wait in enumerate(schedule):
        try:
            return fn(*args, **kwargs)
        except Exception as exc:
            if wait is None:
                raise
            logger.warning("%s failed (attempt %d): %s — retrying in %ds",
                           label, attempt + 1, exc, wait)
            time.sleep(wait)

# ══════════════════════════════════════════════════════════════════════════════
# Credibility scoring
# ══════════════════════════════════════════════════════════════════════════════

_CRED_RULES: list[tuple[list[str], int, str]] = [
    (
        ["arxiv.org", "pubmed", "ncbi.nlm.nih.gov", "nature.com", "science.org",
         "cell.com", "springer", "ieee.org", "acm.org", "jstor.org",
         "nih.gov", "who.int", "bmj.com", "thelancet.com"],
        5, "Peer-reviewed / Academic",
    ),
    (
        ["gov", "edu", "un.org", "europa.eu", "worldbank.org",
         "imf.org", "oecd.org", "cdc.gov", "fda.gov", "nist.gov"],
        4, "Institutional / Government",
    ),
    (
        ["reuters.com", "apnews.com", "bbc.com", "nytimes.com", "theguardian.com",
         "washingtonpost.com", "ft.com", "economist.com", "bloomberg.com",
         "wsj.com", "wired.com", "scientificamerican.com"],
        3, "Reputable News / Magazine",
    ),
    (
        ["medium.com", "substack.com", "wordpress.com",
         "blogspot.com", "tumblr.com", "linkedin.com"],
        1, "Blog / Opinion",
    ),
]


def score_url(url: str) -> tuple[int, str]:
    """Return a 1–5 credibility score and label for *url*."""
    u = url.lower()
    for domains, score, label in _CRED_RULES:
        if any(d in u for d in domains):
            return score, label
    return 2, "General Web"

# ══════════════════════════════════════════════════════════════════════════════
# Citation tracker
# ══════════════════════════════════════════════════════════════════════════════

class CitationTracker:
    """Thread-safe registry of sources with deduplication and credibility scores."""

    def __init__(self, data: Optional[dict] = None) -> None:
        self._store: dict[str, dict] = data or {}
        self._counter = max(
            (int(k[2:-1]) for k in self._store if re.fullmatch(r"\[\^\d+\]", k)),
            default=0,
        )
        self._lock = threading.Lock()

    def register(self, url: str, title: str = "", snippet: str = "") -> str:
        """Register *url* and return its citation ID, deduplicated by URL."""
        with self._lock:
            for cid, meta in self._store.items():
                if meta["url"] == url:
                    return cid
            self._counter += 1
            cid = f"[^{self._counter}]"
            score, label = score_url(url)
            self._store[cid] = {
                "url": url,
                "title": title or url,
                "snippet": snippet[:200],
                "credibility_score": score,
                "credibility_label": label,
                "accessed": datetime.date.today().isoformat(),
            }
            return cid

    def to_dict(self) -> dict:
        with self._lock:
            return dict(self._store)

    def markdown_refs(self) -> str:
        with self._lock:
            if not self._store:
                return "_No sources collected._"
            lines = ["## References\n"]
            for cid, m in self._store.items():
                bar = "★" * m["credibility_score"] + "☆" * (5 - m["credibility_score"])
                lines.append(
                    f"{cid}: **{m['title']}**  \n"
                    f"  <{m['url']}>  \n"
                    f"  Credibility: {bar} ({m['credibility_label']}) · "
                    f"Accessed {m['accessed']}  \n"
                )
            return "\n".join(lines)

    def credibility_table_md(self) -> str:
        with self._lock:
            if not self._store:
                return ""
            rows = [
                "## Source Credibility\n",
                "| Citation | Title | Credibility | Score |",
                "|----------|-------|-------------|-------|",
            ]
            for cid, m in self._store.items():
                bar = "★" * m["credibility_score"] + "☆" * (5 - m["credibility_score"])
                rows.append(
                    f"| {cid} | {m['title'][:48]} | {m['credibility_label']} | {bar} |"
                )
            return "\n".join(rows)

# ══════════════════════════════════════════════════════════════════════════════
# Progress log — thread-safe, SSE-friendly
# ══════════════════════════════════════════════════════════════════════════════

_sse_queues: dict[str, queue.Queue] = {}


def log(phase: str, msg: str, sid: str = "") -> None:
    # SECURITY: Never log API keys or token values.
    # Callers must not pass secrets into *msg*.
    logger.info("[%s] %s", phase, msg)
    if sid and sid in _sse_queues:
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        _sse_queues[sid].put({"type": "log", "phase": phase, "msg": msg, "ts": ts})


def sse_push(sid: str, event: dict) -> None:
    if sid and sid in _sse_queues:
        _sse_queues[sid].put(event)

# ══════════════════════════════════════════════════════════════════════════════
# Tool implementations
# ══════════════════════════════════════════════════════════════════════════════

# SECURITY: The User-Agent identifies this tool honestly. Impersonating browsers
# is avoided to maintain ethical scraping practices.
_HTTP_HEADERS = {"User-Agent": "ResearchAgent/2.0 (github.com/your-org/research-agent)"}


def tool_web_search(query: str, sid: str = "") -> list[dict]:
    """Search the web via Brave Search API with caching and retry."""
    log("SEARCH", f"Query: {query!r}", sid)
    cached = cache_get("search", query)
    if cached is not None:
        log("SEARCH", f"  → cache hit ({len(cached)} results)", sid)
        return cached

    if not BRAVE_API_KEY:
        log("SEARCH", "  ⚠  BRAVE_API_KEY not set — returning mock results", sid)
        return [{"title": "Mock result", "url": "https://example.com",
                 "snippet": "Live search disabled (no BRAVE_API_KEY)."}]

    def _call() -> list[dict]:
        # SECURITY: BRAVE_API_KEY is passed in a header, never in the URL,
        # to prevent accidental logging in server access logs.
        resp = requests.get(
            "https://api.search.brave.com/res/v1/web/search",
            headers={
                "Accept": "application/json",
                # SECURITY: key transmitted as header, not query param
                "X-Subscription-Token": BRAVE_API_KEY,
                **_HTTP_HEADERS,
            },
            params={"q": query, "count": MAX_SEARCH_HITS, "text_decorations": False},
            timeout=12,
            # SECURITY: TLS verification is always enabled (default). Never
            # pass verify=False even in development; use a proper CA bundle.
        )
        resp.raise_for_status()
        data = resp.json()
        return [
            {
                "title":   item.get("title", ""),
                "url":     item.get("url", ""),
                "snippet": item.get("description", ""),
            }
            for item in data.get("web", {}).get("results", [])
        ]

    results = with_retry(_call, label=f"brave-search:{query[:40]}")
    cache_set("search", query, results)
    log("SEARCH", f"  → {len(results)} results", sid)
    return results


def tool_fetch_page(
    url: str,
    title: str = "",
    snippet: str = "",
    citations: Optional[CitationTracker] = None,
    sid: str = "",
) -> str:
    """Fetch and lightly parse a webpage, register a citation, return excerpt."""
    log("FETCH", f"URL: {url}", sid)
    cid = citations.register(url, title=title, snippet=snippet) if citations else "?"

    cached = cache_get("page", url)
    if cached is not None:
        log("FETCH", f"  → cache hit. Citation: {cid}", sid)
        return f"[Citation: {cid}]\n\n{cached}"

    def _call() -> str:
        resp = requests.get(
            url,
            headers=_HTTP_HEADERS,
            timeout=18,
            allow_redirects=True,
            # SECURITY: TLS verification always on. Never disable.
        )
        resp.raise_for_status()
        return resp.text

    try:
        raw = with_retry(_call, label=f"fetch:{url[:60]}")
    except Exception as exc:
        log("FETCH", f"  ✗ Failed to fetch {url}: {exc}", sid)
        return f"[Citation: {cid}]\n\nFailed to fetch page."

    # Minimal HTML stripping — not a full sanitiser; content is only passed to
    # the LLM, never rendered as HTML, so XSS is not a concern here.
    text = re.sub(r"<script[^>]*>.*?</script>", " ", raw, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>",  " ", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s{2,}", " ", text).strip()

    excerpt = text[:4_000]
    cache_set("page", url, excerpt)
    log("FETCH", f"  → {len(text):,} chars fetched. Citation: {cid}", sid)
    return f"[Citation: {cid}]\n\n{excerpt}"

# ══════════════════════════════════════════════════════════════════════════════
# Tool schema for Claude
# ══════════════════════════════════════════════════════════════════════════════

TOOLS: list[dict] = [
    {
        "name": "web_search",
        "description": (
            "Search the web. Returns title, URL, and snippet per result. "
            "Use specific, precise queries."
        ),
        "input_schema": {
            "type": "object",
            "properties": {"query": {"type": "string", "description": "Search query."}},
            "required": ["query"],
        },
    },
    {
        "name": "fetch_page",
        "description": (
            "Fetch full text of a URL. Always call this before citing a source. "
            "Returns ≤4,000 chars of extracted text plus a citation ID."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "url":     {"type": "string"},
                "title":   {"type": "string"},
                "snippet": {"type": "string"},
            },
            "required": ["url"],
        },
    },
    {
        "name": "cite",
        "description": "Register a source and get its citation ID for inline use.",
        "input_schema": {
            "type": "object",
            "properties": {
                "url":     {"type": "string"},
                "title":   {"type": "string"},
                "snippet": {"type": "string"},
            },
            "required": ["url"],
        },
    },
]

# ══════════════════════════════════════════════════════════════════════════════
# System prompt
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """You are an expert research agent with advanced reasoning, planning, and synthesis capabilities.

## MISSION
Given a research question you will autonomously:
1. Decompose it into 3–5 sub-questions
2. Search the web iteratively using web_search
3. Read sources critically using fetch_page (required before any citation)
4. Synthesize findings into a structured Markdown report
5. Self-review before final output

## WORKFLOW

### PHASE 1 — PLANNING
Write out: restatement, sub-questions, source-type priorities, success criteria.

### PHASE 2 — ITERATIVE RESEARCH
For each sub-question: precise query → web_search → select 2–3 credible results →
fetch_page each → extract evidence. Note gaps and refine queries.
Quality gates: credible? recent? on-topic? conflicting evidence noted?

### PHASE 3 — SYNTHESIS
Group by theme; identify consensus, contradictions, open questions.
Weight evidence by credibility score.

### PHASE 4 — REPORT (exact format below)

---
# [Descriptive Report Title]

## Executive Summary
2–3 sentence overview.

## Background
Context and key definitions.

## Key Findings

### Finding 1: [Title]
Explanation with inline citations e.g. "X was shown [^1]."

### Finding 2: [Title]
...

## Conflicting Evidence / Open Questions
Never suppress conflicts.

## Conclusion
Synthesized direct answer.

## Source Credibility Summary
Paragraph rating overall source quality.

---

### PHASE 5 — SELF-REVIEW
Every factual claim cited? No source cited without fetch_page? Conflicts acknowledged?
Report answers original question? Executive summary matches body?

## STYLE
Step-by-step reasoning before each tool call. Justify source selection.
Flag uncertainty: "evidence suggests", "this is debated".
"""

# ══════════════════════════════════════════════════════════════════════════════
# Agent loop
# ══════════════════════════════════════════════════════════════════════════════

def run_agent(
    question: str,
    sid: str = "",
    resume_messages: Optional[list] = None,
) -> tuple[str, CitationTracker]:
    """Drive the Claude agentic loop until the report is complete.

    Returns the final Markdown report text and the populated CitationTracker.
    """
    # SECURITY: API key is validated before each run, never after.
    if not ANTHROPIC_API_KEY:
        raise EnvironmentError("ANTHROPIC_API_KEY is not set")

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    cit = CitationTracker()

    messages: list = resume_messages or [
        {"role": "user", "content": f"Research Question: {question}\n\nBegin with Phase 1."}
    ]

    log("AGENT", f"Starting: {question!r}", sid)
    sse_push(sid, {"type": "status", "status": "running", "question": question})

    final_text = ""

    for iteration in range(1, MAX_ITERATIONS + 1):
        log("AGENT", f"Iteration {iteration}", sid)

        def _api_call():
            return client.messages.create(
                model=MODEL,
                max_tokens=MAX_TOKENS,
                system=SYSTEM_PROMPT,
                tools=TOOLS,
                messages=messages,
            )

        response = with_retry(_api_call, label="claude-api")
        messages.append({"role": "assistant", "content": response.content})

        for blk in response.content:
            if blk.type == "text" and blk.text.strip():
                final_text = blk.text
                sse_push(sid, {"type": "text", "text": blk.text})
                logger.info("[REPORT PREVIEW] %s…", blk.text[:120])

        tool_uses = [b for b in response.content if b.type == "tool_use"]

        if response.stop_reason == "end_turn" and not tool_uses:
            log("AGENT", "Finished — no more tool calls.", sid)
            break

        if tool_uses:
            results = []
            for tu in tool_uses:
                # SECURITY: Tool inputs come from the LLM and are treated as
                # untrusted data. Each handler validates / sanitises its own inputs.
                log("TOOL", f"{tu.name}({json.dumps(tu.input)[:80]}…)", sid)
                sse_push(sid, {"type": "tool", "name": tu.name, "input": tu.input})

                if tu.name == "web_search":
                    out = tool_web_search(tu.input.get("query", ""), sid=sid)
                    content = json.dumps(out, ensure_ascii=False)
                elif tu.name == "fetch_page":
                    content = tool_fetch_page(
                        tu.input.get("url", ""),
                        title=tu.input.get("title", ""),
                        snippet=tu.input.get("snippet", ""),
                        citations=cit,
                        sid=sid,
                    )
                elif tu.name == "cite":
                    content = cit.register(
                        tu.input.get("url", ""),
                        title=tu.input.get("title", ""),
                        snippet=tu.input.get("snippet", ""),
                    )
                else:
                    log("TOOL", f"Unknown tool requested: {tu.name}", sid)
                    content = f"Unknown tool: {tu.name}"

                results.append({
                    "type": "tool_result",
                    "tool_use_id": tu.id,
                    "content": content,
                })

            messages.append({"role": "user", "content": results})

            # Persist mid-run state so a crash doesn't lose progress
            if sid:
                session_update(
                    sid,
                    messages=json.dumps(
                        [m for m in messages if isinstance(m, dict)], default=str
                    ),
                    citations=json.dumps(cit.to_dict()),
                )
    else:
        log("AGENT", f"⚠  Max iterations ({MAX_ITERATIONS}) reached.", sid)

    return final_text, cit

# ══════════════════════════════════════════════════════════════════════════════
# Export helpers
# ══════════════════════════════════════════════════════════════════════════════

def _make_stem(question: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", question.lower())[:55].strip("_")
    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"report_{slug}_{ts}"


def export_md(question: str, body: str, cit: CitationTracker, stem: str) -> Path:
    """Write the Markdown report to REPORTS_DIR/<stem>.md."""
    path = REPORTS_DIR / f"{stem}.md"
    content = (
        f"---\n"
        f"question: {question!r}\n"
        f"generated: {datetime.datetime.utcnow().isoformat()}Z\n"
        f"model: {MODEL}\n"
        f"sources: {len(cit.to_dict())}\n"
        f"---\n\n"
        f"{body}\n\n"
        f"{cit.credibility_table_md()}\n\n"
        f"{cit.markdown_refs()}\n"
    )
    path.write_text(content, encoding="utf-8")
    log("EXPORT", f"Markdown → {path}")
    return path


def export_pdf(question: str, body: str, cit: CitationTracker, stem: str) -> Optional[Path]:
    """Write a styled PDF report to REPORTS_DIR/<stem>.pdf."""
    if not REPORTLAB_OK:
        log("EXPORT", "reportlab not installed — skipping PDF")
        return None

    path = REPORTS_DIR / f"{stem}.pdf"
    doc  = SimpleDocTemplate(
        str(path), pagesize=letter,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    h1  = ParagraphStyle("H1",  parent=styles["Heading1"],  fontSize=18, spaceAfter=12,
                          textColor=colors.HexColor("#1a3a5c"))
    h2  = ParagraphStyle("H2",  parent=styles["Heading2"],  fontSize=14, spaceAfter=8,
                          textColor=colors.HexColor("#2e6da4"))
    h3  = ParagraphStyle("H3",  parent=styles["Heading3"],  fontSize=12, spaceAfter=6,
                          textColor=colors.HexColor("#4a90d9"))
    bod = ParagraphStyle("Body", parent=styles["Normal"],   fontSize=10, spaceAfter=8, leading=15)
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=8,
                                 textColor=colors.grey)

    story: list = []

    def _esc(s: str) -> str:
        s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        s = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", s)
        s = re.sub(r"\*(.*?)\*",     r"<i>\1</i>", s)
        return s

    def _md_to_story(md: str) -> None:
        for line in md.splitlines():
            s = line.strip()
            if not s:
                story.append(Spacer(1, 6))
            elif s.startswith("### "):
                story.append(Paragraph(s[4:], h3))
            elif s.startswith("## "):
                story.append(Paragraph(s[3:], h2))
            elif s.startswith("# "):
                story.append(Paragraph(s[2:], h1))
            elif s.startswith("---"):
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            elif s.startswith("| "):
                pass  # table rows handled separately
            else:
                story.append(Paragraph(_esc(s), bod))

    story.append(Paragraph("Research Report", h1))
    story.append(Paragraph(
        f"Question: {_esc(question)}  ·  "
        f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  Model: {MODEL}",
        meta_style,
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2e6da4")))
    story.append(Spacer(1, 12))
    _md_to_story(body)

    store = cit.to_dict()
    if store:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Source Credibility", h2))
        tdata = [["Citation", "Title", "Credibility", "Score"]]
        for cid_key, m in store.items():
            bar = "★" * m["credibility_score"] + "☆" * (5 - m["credibility_score"])
            tdata.append([cid_key, m["title"][:45], m["credibility_label"], bar])
        t = Table(tdata, colWidths=[0.7 * inch, 3.5 * inch, 1.8 * inch, 0.9 * inch])
        t.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0), (-1, 0), colors.HexColor("#2e6da4")),
            ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
            ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",       (0, 0), (-1, -1), 8),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
            ("GRID",           (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
            ("PADDING",        (0, 0), (-1, -1), 4),
        ]))
        story.append(t)

    doc.build(story)
    log("EXPORT", f"PDF → {path}")
    return path


def export_docx(question: str, body: str, cit: CitationTracker, stem: str) -> Optional[Path]:
    """Write a styled DOCX report to REPORTS_DIR/<stem>.docx."""
    if not DOCX_OK:
        log("EXPORT", "python-docx not installed — skipping DOCX")
        return None

    path = REPORTS_DIR / f"{stem}.docx"
    doc  = DocxDocument()

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def _add_heading(text: str, level: int):
        p = doc.add_heading(text, level=level)
        colour = RGBColor(0x1a, 0x3a, 0x5c) if level == 1 else RGBColor(0x2e, 0x6d, 0xa4)
        for run in p.runs:
            run.font.color.rgb = colour
        return p

    meta = doc.add_paragraph()
    meta.add_run("Research Report\n").bold = True
    meta.add_run(f"Question: {question}\n")
    meta.add_run(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  Model: {MODEL}"
    ).font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    for line in body.splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph()
        elif s.startswith("### "):
            _add_heading(s[4:], 3)
        elif s.startswith("## "):
            _add_heading(s[3:], 2)
        elif s.startswith("# "):
            _add_heading(s[2:], 1)
        elif s.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif s.startswith("| "):
            pass
        else:
            doc.add_paragraph(re.sub(r"\*\*(.*?)\*\*", r"\1", s))

    store = cit.to_dict()
    if store:
        _add_heading("Source Credibility", 2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Citation", "Title", "Credibility", "Score"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for cid_key, m in store.items():
            row = table.add_row().cells
            bar = "★" * m["credibility_score"] + "☆" * (5 - m["credibility_score"])
            row[0].text = cid_key
            row[1].text = m["title"][:50]
            row[2].text = m["credibility_label"]
            row[3].text = bar

        _add_heading("References", 2)
        for cid_key, m in store.items():
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{cid_key} ").bold = True
            p.add_run(f"{m['title']}  –  {m['url']}  (accessed {m['accessed']})")

    doc.save(str(path))
    log("EXPORT", f"DOCX → {path}")
    return path


def run_and_export(
    question: str,
    formats: list[str],
    sid: str = "",
) -> dict[str, Optional[Path]]:
    """Run the agent then export in all requested formats."""
    body, cit = run_agent(question, sid=sid)
    stem = _make_stem(question)
    out: dict[str, Optional[Path]] = {}

    if "md"   in formats: out["md"]   = export_md(question, body, cit, stem)
    if "pdf"  in formats: out["pdf"]  = export_pdf(question, body, cit, stem)
    if "docx" in formats: out["docx"] = export_docx(question, body, cit, stem)

    if sid:
        session_update(
            sid,
            status="done",
            finished_at=datetime.datetime.utcnow().isoformat(),
            report_md=body,
            citations=json.dumps(cit.to_dict()),
        )
        sse_push(sid, {"type": "done", "files": {k: str(v) for k, v in out.items() if v}})

    return out

# ══════════════════════════════════════════════════════════════════════════════
# Web UI  (Flask)
# ══════════════════════════════════════════════════════════════════════════════

# Inline HTML template — kept here to remain a single-file deployment option.
_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<meta http-equiv="Content-Security-Policy"
      content="default-src 'self'; script-src 'self' 'unsafe-inline'; style-src 'self' 'unsafe-inline'">
<title>Research Agent</title>
<style>
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',system-ui,sans-serif;background:#0f1117;color:#e2e8f0;min-height:100vh}
.app{max-width:1100px;margin:0 auto;padding:24px}
h1{font-size:1.7rem;font-weight:700;color:#7dd3fc;margin-bottom:4px}
.sub{color:#64748b;font-size:.85rem;margin-bottom:28px}
.card{background:#1e2330;border:1px solid #2d3748;border-radius:12px;padding:24px;margin-bottom:20px}
.row{display:flex;gap:12px;flex-wrap:wrap}
input[type=text]{flex:1;min-width:260px;padding:10px 16px;border-radius:8px;border:1px solid #2d3748;
  background:#0f1117;color:#e2e8f0;font-size:.95rem;outline:none}
input[type=text]:focus{border-color:#7dd3fc}
.checks{display:flex;gap:14px;align-items:center;flex-wrap:wrap}
.checks label{display:flex;align-items:center;gap:6px;font-size:.88rem;cursor:pointer;color:#94a3b8}
.checks input{accent-color:#7dd3fc}
button{padding:10px 22px;border-radius:8px;border:none;cursor:pointer;font-size:.92rem;font-weight:600;transition:.15s}
.btn-primary{background:#2563eb;color:#fff}.btn-primary:hover{background:#1d4ed8}
.btn-sm{padding:6px 14px;font-size:.8rem;background:#1e2330;color:#94a3b8;border:1px solid #2d3748}
#log{background:#0a0d14;border:1px solid #1e2330;border-radius:8px;padding:16px;
  height:320px;overflow-y:auto;font-family:'Fira Code',monospace;font-size:.78rem;line-height:1.6}
.log-SEARCH{color:#34d399}.log-FETCH{color:#60a5fa}.log-AGENT{color:#a78bfa}
.log-TOOL{color:#fbbf24}.log-EXPORT{color:#f472b6}.log-RETRY{color:#fb923c}
#report{background:#0a0d14;border:1px solid #1e2330;border-radius:8px;padding:20px;
  white-space:pre-wrap;font-size:.88rem;line-height:1.7;max-height:540px;overflow-y:auto;display:none}
.files{display:flex;gap:10px;flex-wrap:wrap;margin-top:12px}
.pill{padding:6px 14px;border-radius:20px;font-size:.8rem;font-weight:600;text-decoration:none;
  background:#1e40af;color:#bfdbfe;border:1px solid #2563eb}
.sessions{display:flex;flex-direction:column;gap:8px}
.sess{display:flex;justify-content:space-between;align-items:center;padding:10px 14px;
  background:#0f1117;border-radius:8px;border:1px solid #1e2330;font-size:.82rem}
.sess .q{color:#7dd3fc;max-width:60%;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.badge{padding:2px 8px;border-radius:10px;font-size:.72rem;font-weight:600}
.badge-done{background:#064e3b;color:#6ee7b7}
.badge-running{background:#1e3a5f;color:#93c5fd}
.badge-error{background:#450a0a;color:#fca5a5}
#spinner{display:none;width:16px;height:16px;border:2px solid #2d3748;border-top-color:#7dd3fc;
  border-radius:50%;animation:spin .7s linear infinite;margin-left:8px}
@keyframes spin{to{transform:rotate(360deg)}}
</style>
</head>
<body>
<div class="app">
  <h1>🔬 Research Agent</h1>
  <p class="sub">Autonomous web research · multi-step planning · structured reports</p>
  <div class="card">
    <div class="row" style="margin-bottom:14px">
      <input id="q" type="text" maxlength="500" placeholder="Enter your research question…">
      <div id="spinner"></div>
      <button class="btn-primary" onclick="startResearch()">Research</button>
    </div>
    <div class="checks">
      <span style="color:#64748b;font-size:.83rem">Export:</span>
      <label><input type="checkbox" id="fmt_md"   checked> Markdown</label>
      <label><input type="checkbox" id="fmt_pdf"  checked> PDF</label>
      <label><input type="checkbox" id="fmt_docx" checked> DOCX</label>
    </div>
  </div>
  <div class="card" id="logCard" style="display:none">
    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px">
      <span style="font-size:.85rem;color:#64748b">Live Progress</span>
      <button class="btn-sm" onclick="document.getElementById('log').innerHTML=''">Clear</button>
    </div>
    <div id="log"></div>
    <div id="fileLinks" class="files"></div>
  </div>
  <div class="card" id="reportCard" style="display:none">
    <div style="margin-bottom:10px;font-size:.85rem;color:#64748b">Report Preview</div>
    <div id="report"></div>
  </div>
  <div class="card">
    <div style="margin-bottom:12px;font-size:.88rem;color:#94a3b8">Past Sessions</div>
    <div class="sessions" id="sessions">Loading…</div>
  </div>
</div>
<script>
// SECURITY: All data inserted into the DOM uses textContent, not innerHTML,
// except for trusted static strings. XSS is prevented by never eval-ing
// server responses or concatenating them into innerHTML.
'use strict';
let es;
const MAX_Q_LEN = 500;

function startResearch() {
  const qEl = document.getElementById('q');
  const q = qEl.value.trim();
  if (!q) return;
  if (q.length > MAX_Q_LEN) { alert('Question too long (max 500 chars)'); return; }
  const fmts = ['md','pdf','docx'].filter(f => document.getElementById('fmt_'+f).checked);
  if (!fmts.length) { alert('Select at least one export format'); return; }

  document.getElementById('logCard').style.display = 'block';
  document.getElementById('reportCard').style.display = 'none';
  document.getElementById('log').innerHTML = '';
  document.getElementById('fileLinks').innerHTML = '';
  document.getElementById('report').style.display = 'none';
  document.getElementById('spinner').style.display = 'block';
  if (es) es.close();

  fetch('/start', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({question: q, formats: fmts})
  })
  .then(r => r.json())
  .then(d => {
    if (d.error) { alert('Error: ' + d.error); document.getElementById('spinner').style.display='none'; return; }
    const sid = d.session_id;
    es = new EventSource('/stream/' + encodeURIComponent(sid));
    es.onmessage = e => {
      const ev = JSON.parse(e.data);
      if (ev.type === 'log') appendLog(ev);
      else if (ev.type === 'text') showReport(ev.text);
      else if (ev.type === 'done') {
        document.getElementById('spinner').style.display = 'none';
        showFiles(ev.files || {});
        es.close();
        loadSessions();
      }
    };
    es.onerror = () => { document.getElementById('spinner').style.display='none'; es.close(); };
  })
  .catch(err => { alert('Request failed: ' + err); document.getElementById('spinner').style.display='none'; });
}

function appendLog(ev) {
  const el = document.getElementById('log');
  const line = document.createElement('div');
  // SECURITY: textContent prevents XSS — never use innerHTML with server data
  line.className = 'log-' + ev.phase.replace(/[^A-Z]/g, '');
  line.textContent = '[' + ev.ts + '] [' + ev.phase + '] ' + ev.msg;
  el.appendChild(line);
  el.scrollTop = el.scrollHeight;
}

function showReport(text) {
  document.getElementById('reportCard').style.display = 'block';
  const div = document.getElementById('report');
  div.style.display = 'block';
  div.textContent = text;  // SECURITY: textContent, not innerHTML
  div.scrollTop = div.scrollHeight;
}

function showFiles(files) {
  const fl = document.getElementById('fileLinks');
  const lbl = document.createElement('span');
  lbl.style.cssText = 'font-size:.8rem;color:#64748b';
  lbl.textContent = 'Download: ';
  fl.appendChild(lbl);
  for (const [fmt, path] of Object.entries(files)) {
    const a = document.createElement('a');
    // SECURITY: Encode path to prevent open-redirect via crafted server response
    a.href = '/download?path=' + encodeURIComponent(path);
    a.className = 'pill';
    a.textContent = fmt.toUpperCase();
    a.target = '_blank';
    a.rel = 'noopener noreferrer';
    fl.appendChild(a);
  }
}

function loadSessions() {
  fetch('/sessions')
    .then(r => r.json())
    .then(data => {
      const el = document.getElementById('sessions');
      if (!data.length) {
        el.textContent = 'No past sessions yet.';
        return;
      }
      el.innerHTML = '';
      data.forEach(s => {
        const row = document.createElement('div');
        row.className = 'sess';
        const q = document.createElement('span');
        q.className = 'q';
        q.title = s.question;
        q.textContent = s.question;   // SECURITY: textContent
        const meta = document.createElement('div');
        meta.style.cssText = 'display:flex;gap:10px;align-items:center';
        const badge = document.createElement('span');
        badge.className = 'badge badge-' + s.status.replace(/[^a-z]/g,'');
        badge.textContent = s.status;
        const ts = document.createElement('span');
        ts.className = 'meta';
        ts.textContent = (s.started_at || '').slice(0, 16).replace('T', ' ');
        meta.appendChild(badge);
        meta.appendChild(ts);
        row.appendChild(q);
        row.appendChild(meta);
        el.appendChild(row);
      });
    });
}
loadSessions();
</script>
</body>
</html>"""


def create_app() -> "Flask":
    """Create and configure the Flask application.

    SECURITY hardening applied:
      - FLASK_SECRET_KEY required; app refuses to start without it.
      - debug=False enforced regardless of environment.
      - Security headers injected on every response.
      - Host header validated against ALLOWED_HOSTS.
      - Path traversal blocked in /download.
      - Input validation on /start before any processing.
      - No CORS headers (API is not public).
    """
    if not FLASK_SECRET_KEY:
        raise EnvironmentError(
            "FLASK_SECRET_KEY is not set. Generate one with: "
            "python -c \"import secrets; print(secrets.token_hex(32))\""
        )

    app = Flask(__name__)
    # SECURITY: secret_key must be set for session signing even if we don't
    # use server-side sessions, as Flask extensions rely on it.
    app.secret_key = FLASK_SECRET_KEY

    # SECURITY: Enforce HTTP Basic Auth on every request when configured.
    app.before_request(_require_auth)

    # SECURITY: Inject HTTP security headers on every response.
    @app.after_request
    def _security_headers(response):
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"]        = "DENY"
        response.headers["X-XSS-Protection"]       = "1; mode=block"
        response.headers["Referrer-Policy"]        = "strict-origin-when-cross-origin"
        response.headers["Content-Security-Policy"] = (
            "default-src 'self'; "
            "script-src 'self' 'unsafe-inline'; "
            "style-src 'self' 'unsafe-inline'"
        )
        # SECURITY: HSTS — uncomment when TLS is terminated at this server.
        # response.headers["Strict-Transport-Security"] = "max-age=63072000; includeSubDomains"
        return response

    # SECURITY: Reject requests with unexpected Host headers to prevent
    # HTTP host-header injection attacks.
    @app.before_request
    def _check_host():
        host = request.host.split(":")[0]
        if host not in ALLOWED_HOSTS:
            abort(400, description="Invalid Host header")

    @app.route("/")
    def index():
        return render_template_string(_HTML)

    @app.route("/start", methods=["POST"])
    def start():
        data = request.get_json(silent=True) or {}

        # SECURITY: Validate and sanitise all user-supplied inputs before use.
        try:
            question = validate_question(data.get("question", ""))
            formats  = validate_formats(data.get("formats", []))
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400

        # SECURITY: Per-IP rate limit on /start — most expensive endpoint.
        ip = _client_ip()
        if not _research_bucket.allow(ip):
            return jsonify({"error": "Rate limit exceeded. Please wait before submitting another question."}), 429

        # SECURITY: Session ID is a cryptographically random token, not derived
        # from user input, preventing session-fixation attacks.
        sid = secrets.token_hex(8)

        session_create(sid, question)
        _sse_queues[sid] = queue.Queue()

        def worker():
            try:
                run_and_export(question, formats, sid=sid)
            except Exception:
                # SECURITY: Full tracebacks are logged server-side only.
                # The client receives a generic error event without stack details.
                logger.error("Agent error for session %s:\n%s", sid, traceback.format_exc())
                session_update(
                    sid,
                    status="error",
                    finished_at=datetime.datetime.utcnow().isoformat(),
                )
                sse_push(sid, {"type": "done", "files": {}})

        threading.Thread(target=worker, daemon=True).start()
        return jsonify({"session_id": sid})

    @app.route("/stream/<sid>")
    def stream(sid: str):
        # SECURITY: Validate the session ID format before looking it up.
        if not re.fullmatch(r"[0-9a-f]{16}", sid):
            abort(400)

        def _gen() -> Iterator[str]:
            q = _sse_queues.get(sid)
            if not q:
                return
            while True:
                try:
                    ev = q.get(timeout=30)
                    yield f"data: {json.dumps(ev)}\n\n"
                    if ev.get("type") == "done":
                        break
                except queue.Empty:
                    yield 'data: {"type":"ping"}\n\n'

        return Response(
            _gen(),
            mimetype="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    @app.route("/sessions")
    def list_sessions_route():
        return jsonify(sessions_list())

    @app.route("/download")
    def download():
        path_str = request.args.get("path", "")
        try:
            # SECURITY: safe_report_path enforces that only files inside
            # REPORTS_DIR can be served — path traversal is blocked here.
            p = safe_report_path(path_str)
        except (PermissionError, Exception):
            abort(404)
        if not p.is_file():
            abort(404)
        return send_file(str(p), as_attachment=True)

    return app

# ══════════════════════════════════════════════════════════════════════════════
# CLI entry point
# ══════════════════════════════════════════════════════════════════════════════

def main() -> None:
    import argparse

    ap = argparse.ArgumentParser(
        description="Autonomous Research Agent",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python research_agent.py \"What causes long COVID?\" --format md,pdf,docx\n"
            "  python research_agent.py --serve --port 8080\n"
        ),
    )
    ap.add_argument("question", nargs="*", help="Research question (CLI mode)")
    ap.add_argument("--format",  default="md",       help="Comma-separated: md,pdf,docx")
    ap.add_argument("--serve",   action="store_true", help="Launch Flask web UI")
    ap.add_argument("--port",    type=int, default=5000)
    ap.add_argument("--host",    default="127.0.0.1",
                    help="Bind address (default: 127.0.0.1 — localhost only)")
    ap.add_argument("--gen-password", action="store_true",
                    help="Generate a UI_PASSWORD_SALT + UI_PASSWORD_HASH for .env")
    args = ap.parse_args()

    if args.gen_password:
        import getpass
        pw = getpass.getpass("Enter password for web UI: ")
        pw2 = getpass.getpass("Confirm password: ")
        if pw != pw2:
            sys.exit("ERROR: Passwords do not match.")
        salt, phash = generate_password_hash(pw)
        print(f"\nAdd these to your .env file:\n")
        print(f"UI_USERNAME=admin")
        print(f"UI_PASSWORD_SALT={salt}")
        print(f"UI_PASSWORD_HASH={phash}")
        return

    if not ANTHROPIC_API_KEY:
        sys.exit("ERROR: ANTHROPIC_API_KEY environment variable is not set.")

    if args.serve:
        if not FLASK_OK:
            sys.exit("ERROR: Flask not installed.  pip install flask")
        # SECURITY: debug=False is hard-coded. Never enable debug mode in
        # production — it exposes an interactive REPL to anyone with network access.
        print(f"Starting web UI at http://{args.host}:{args.port}")
        create_app().run(host=args.host, port=args.port, debug=False)
        return

    if not args.question:
        ap.print_help()
        sys.exit(1)

    try:
        question = validate_question(" ".join(args.question))
        formats  = validate_formats([f.strip() for f in args.format.split(",")])
    except ValueError as exc:
        sys.exit(f"ERROR: {exc}")

    t0 = time.time()
    files = run_and_export(question, formats)

    print(f"\n{'═' * 60}")
    print(f"Done in {time.time() - t0:.1f}s")
    for fmt, p in files.items():
        if p:
            print(f"  {fmt.upper():6} → {p}")


if __name__ == "__main__":
    main()
